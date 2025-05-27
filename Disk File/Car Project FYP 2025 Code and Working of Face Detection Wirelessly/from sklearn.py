from sklearn.neighbors import KNeighborsClassifier
import cv2
import pickle
import numpy as np
import os
import csv
import time
from datetime import datetime
from win32com.client import Dispatch
from docx import Document

def speak(str1):
    speak = Dispatch(("SAPI.SpVoice"))
    speak.Speak(str1)

# Replace this with your ESP32 camera stream URL
stream_url = "http://192.168.204.10:81/stream"  # Adjust if necessary"  # Adjust if necessary

# Open the stream with OpenCV
video = cv2.VideoCapture(stream_url)

if not video.isOpened():
    print("Error: Unable to connect to stream.")
    exit()

facedetect = cv2.CascadeClassifier('data/haarcascade_frontalface_default.xml')

with open('data/names.pkl', 'rb') as w:
    LABELS = pickle.load(w)
with open('data/faces_data.pkl', 'rb') as f:
    FACES = pickle.load(f)

print('Shape of Faces matrix --> ', FACES.shape)

knn = KNeighborsClassifier(n_neighbors=5)
knn.fit(FACES, LABELS)

COL_NAMES = ['NAME', 'TIME']

# Create or load the Word document for attendance log
attendance_doc = Document()
attendance_doc.add_heading('Face Detection Attendance Log', 0)

# Initialize dictionary to store detection counts
detection_counts = {}

# Set up video recording (we're using mp4 format for video)
fourcc = cv2.VideoWriter_fourcc(*'XVID')
video_filename = f"Attendance_Video_{datetime.now().strftime('%d-%m-%Y_%H-%M-%S')}.avi"
out_video = cv2.VideoWriter(video_filename, fourcc, 20.0, (640, 480))

while True:
    ret, frame = video.read()

    # Check if the frame is correctly captured
    if not ret:
        print("Failed to grab frame, trying again...")
        time.sleep(0.1)  # Add a small delay before retrying
        continue

    gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
    faces = facedetect.detectMultiScale(gray, 1.3, 5)

    for (x, y, w, h) in faces:
        crop_img = frame[y:y+h, x:x+w, :]
        resized_img = cv2.resize(crop_img, (50, 50)).flatten().reshape(1, -1)
        output = knn.predict(resized_img)
        ts = time.time()
        date = datetime.fromtimestamp(ts).strftime("%d-%m-%Y")
        timestamp = datetime.fromtimestamp(ts).strftime("%H:%M:%S")
        exist = os.path.isfile(f"Attendance/Attendance_{date}.csv")
        
        # Draw face bounding boxes and name
        cv2.rectangle(frame, (x, y), (x+w, y+h), (0, 0, 255), 1)
        cv2.rectangle(frame, (x, y), (x+w, y+h), (50, 50, 255), 2)
        cv2.rectangle(frame, (x, y-40), (x+w, y), (50, 50, 255), -1)
        cv2.putText(frame, str(output[0]), (x, y-15), cv2.FONT_HERSHEY_COMPLEX, 1, (255, 255, 255), 1)
        
        # Attendance info to save
        attendance = [str(output[0]), str(timestamp)]

        # Update detection count
        if output[0] in detection_counts:
            detection_counts[output[0]] += 1
        else:
            detection_counts[output[0]] = 1

        # Update the Word document
        attendance_doc.add_paragraph(f"{output[0]} detected at {timestamp} - Count: {detection_counts[output[0]]}")

    # Display the frame and write it to the video file
    out_video.write(frame)
    cv2.imshow("Frame", frame)

    k = cv2.waitKey(1)

    if k == ord('o'):
        speak("Attendance Taken..")
        time.sleep(5)
        # Log attendance in CSV
        if exist:
            with open(f"Attendance/Attendance_{date}.csv", "+a") as csvfile:
                writer = csv.writer(csvfile)
                writer.writerow(attendance)
            csvfile.close()
        else:
            with open(f"Attendance/Attendance_{date}.csv", "+a") as csvfile:
                writer = csv.writer(csvfile)
                writer.writerow(COL_NAMES)
                writer.writerow(attendance)
            csvfile.close()

    if k == ord('q'):
        break

# Save the Word document at the end
attendance_doc.save(f"Attendance_Log_{datetime.now().strftime('%d-%m-%Y_%H-%M-%S')}.docx")

# Release resources
video.release()
out_video.release()
cv2.destroyAllWindows()
